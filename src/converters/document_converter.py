"""
Conversor principal de documentos Word a PDF.
"""

import os
import time
from pathlib import Path
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from utils import LoggerMixin, ConversionError, PDFGenerationError, log_execution_time
from config.settings import config
from .pdf_generator import PDFGenerator


class DocumentConverter(LoggerMixin):
    """Conversor de documentos Word a PDF."""
    
    def __init__(self):
        super().__init__("DocumentConverter")
        self.pdf_generator = PDFGenerator()
        self.conversion_stats = {
            'total_conversions': 0,
            'successful_conversions': 0,
            'failed_conversions': 0,
            'total_time': 0.0
        }
    
    @log_execution_time
    def convert_document(self, input_path: str, output_path: str) -> bool:
        """
        Convierte un documento Word a PDF.
        
        Args:
            input_path: Ruta del archivo de entrada (.doc/.docx)
            output_path: Ruta del archivo de salida (.pdf)
            
        Returns:
            bool: True si la conversión fue exitosa
            
        Raises:
            ConversionError: Si hay un error en la conversión
        """
        try:
            self.log_info(f"Iniciando conversión: {input_path} -> {output_path}")
            
            # Verificar archivo de entrada
            if not os.path.exists(input_path):
                raise ConversionError(f"Archivo de entrada no encontrado: {input_path}")
            
            # Crear directorio de salida si no existe
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Cargar documento Word
            doc = self._load_document(input_path)
            
            # Extraer contenido del documento
            content = self._extract_content(doc)
            
            # Generar PDF
            success = self.pdf_generator.generate_pdf(content, output_path)
            
            if success:
                self.conversion_stats['successful_conversions'] += 1
                self.log_info(f"Conversión exitosa: {output_path}")
            else:
                self.conversion_stats['failed_conversions'] += 1
                raise PDFGenerationError(f"No se pudo generar el PDF: {output_path}")
            
            return success
            
        except Exception as e:
            self.conversion_stats['failed_conversions'] += 1
            self.log_error(f"Error en conversión de {input_path}: {e}")
            raise ConversionError(f"Error de conversión: {str(e)}", input_path)
    
    def _load_document(self, file_path: str) -> Document:
        """
        Carga un documento Word.
        
        Args:
            file_path: Ruta del archivo
            
        Returns:
            Document: Documento cargado
            
        Raises:
            ConversionError: Si no se puede cargar el documento
        """
        try:
            self.log_debug(f"Cargando documento: {file_path}")
            
            # Verificar extensión
            if not file_path.lower().endswith(('.doc', '.docx')):
                raise ConversionError(f"Formato no soportado: {Path(file_path).suffix}")
            
            # Para archivos .doc, usar antiword para extraer texto
            if file_path.lower().endswith('.doc'):
                return self._load_doc_file(file_path)
            else:
                # Para archivos .docx, usar python-docx
                doc = Document(file_path)
                self.log_debug(f"Documento cargado: {len(doc.paragraphs)} párrafos")
                return doc
            
        except Exception as e:
            raise ConversionError(f"No se pudo cargar el documento: {str(e)}", file_path)
    
    def _load_doc_file(self, file_path: str) -> Document:
        """
        Carga un archivo .doc usando antiword.
        
        Args:
            file_path: Ruta del archivo .doc
            
        Returns:
            Document: Documento cargado como .docx
            
        Raises:
            ConversionError: Si no se puede cargar el documento
        """
        try:
            import subprocess
            import tempfile
            
            self.log_debug(f"Procesando archivo .doc con antiword: {file_path}")
            
            # Extraer texto usando antiword
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore'
            )
            
            if result.returncode != 0:
                raise ConversionError(f"Error ejecutando antiword: {result.stderr}")
            
            text_content = result.stdout
            
            # Crear un nuevo documento .docx con el contenido extraído
            doc = Document()
            
            # Dividir el texto en párrafos y agregarlos al documento
            paragraphs = text_content.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)
            
            self.log_debug(f"Archivo .doc convertido: {len(doc.paragraphs)} párrafos")
            return doc
            
        except Exception as e:
            raise ConversionError(f"Error procesando archivo .doc: {str(e)}", file_path)
    
    def _extract_content(self, doc: Document) -> Dict[str, Any]:
        """
        Extrae el contenido del documento Word.
        
        Args:
            doc: Documento Word cargado
            
        Returns:
            Dict[str, Any]: Contenido extraído estructurado
        """
        content = {
            'title': '',
            'paragraphs': [],
            'tables': [],
            'images': [],
            'styles': {},
            'metadata': {}
        }
        
        try:
            # Extraer título del documento
            if doc.core_properties.title:
                content['title'] = doc.core_properties.title
            elif doc.paragraphs:
                # Usar el primer párrafo como título si no hay título definido
                first_para = doc.paragraphs[0]
                if first_para.text.strip():
                    content['title'] = first_para.text.strip()
            
            # Extraer párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraph_data = self._extract_paragraph(para)
                    content['paragraphs'].append(paragraph_data)
            
            # Extraer tablas
            for table in doc.tables:
                table_data = self._extract_table(table)
                content['tables'].append(table_data)
            
            # Extraer imágenes
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_data = self._extract_image(rel)
                    if image_data:
                        content['images'].append(image_data)
            
            # Extraer metadatos
            content['metadata'] = self._extract_metadata(doc)
            
            self.log_debug(f"Contenido extraído: {len(content['paragraphs'])} párrafos, "
                          f"{len(content['tables'])} tablas, {len(content['images'])} imágenes")
            
            return content
            
        except Exception as e:
            self.log_error(f"Error extrayendo contenido: {e}")
            raise ConversionError(f"Error extrayendo contenido: {str(e)}")
    
    def _extract_paragraph(self, para) -> Dict[str, Any]:
        """
        Extrae información de un párrafo.
        
        Args:
            para: Párrafo del documento
            
        Returns:
            Dict[str, Any]: Datos del párrafo
        """
        paragraph_data = {
            'text': para.text,
            'style': para.style.name if para.style else 'Normal',
            'alignment': str(para.alignment) if para.alignment else 'left',
            'runs': []
        }
        
        # Extraer runs (fragmentos de texto con formato)
        for run in para.runs:
            run_data = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_name': run.font.name if run.font.name else None,
                'color': str(run.font.color.rgb) if run.font.color.rgb else None
            }
            paragraph_data['runs'].append(run_data)
        
        return paragraph_data
    
    def _extract_table(self, table) -> Dict[str, Any]:
        """
        Extrae información de una tabla.
        
        Args:
            table: Tabla del documento
            
        Returns:
            Dict[str, Any]: Datos de la tabla
        """
        table_data = {
            'rows': [],
            'columns': len(table.columns),
            'alignment': str(table.alignment) if table.alignment else 'left'
        }
        
        # Extraer filas
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_data = {
                    'text': cell.text,
                    'paragraphs': []
                }
                
                # Extraer párrafos de la celda
                for para in cell.paragraphs:
                    if para.text.strip():
                        cell_data['paragraphs'].append(self._extract_paragraph(para))
                
                row_data.append(cell_data)
            table_data['rows'].append(row_data)
        
        return table_data
    
    def _extract_image(self, rel) -> Optional[Dict[str, Any]]:
        """
        Extrae información de una imagen.
        
        Args:
            rel: Relación de imagen
            
        Returns:
            Optional[Dict[str, Any]]: Datos de la imagen
        """
        try:
            image_data = {
                'target': rel.target_ref,
                'type': rel.reltype,
                'data': rel.target_part.blob if hasattr(rel, 'target_part') else None
            }
            return image_data
        except Exception as e:
            self.log_warning(f"No se pudo extraer imagen: {e}")
            return None
    
    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """
        Extrae metadatos del documento.
        
        Args:
            doc: Documento Word
            
        Returns:
            Dict[str, Any]: Metadatos del documento
        """
        metadata = {}
        
        try:
            props = doc.core_properties
            metadata.update({
                'title': props.title,
                'author': props.author,
                'subject': props.subject,
                'keywords': props.keywords,
                'comments': props.comments,
                'category': props.category,
                'created': props.created,
                'modified': props.modified,
                'last_modified_by': props.last_modified_by,
                'revision': props.revision
            })
        except Exception as e:
            self.log_warning(f"No se pudieron extraer metadatos: {e}")
        
        return metadata
    
    def get_conversion_stats(self) -> Dict[str, Any]:
        """
        Obtiene estadísticas de conversión.
        
        Returns:
            Dict[str, Any]: Estadísticas de conversión
        """
        total = self.conversion_stats['successful_conversions'] + self.conversion_stats['failed_conversions']
        success_rate = (self.conversion_stats['successful_conversions'] / total * 100) if total > 0 else 0
        
        return {
            **self.conversion_stats,
            'success_rate': success_rate
        }
    
    def reset_stats(self):
        """Reinicia las estadísticas de conversión."""
        self.conversion_stats = {
            'total_conversions': 0,
            'successful_conversions': 0,
            'failed_conversions': 0,
            'total_time': 0.0
        } 