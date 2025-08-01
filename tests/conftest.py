"""
Configuración de pytest y fixtures comunes para tests.
"""

import pytest
import tempfile
import os
from pathlib import Path
from unittest.mock import Mock, patch

# Agregar src al path para imports
import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))


@pytest.fixture
def temp_dir():
    """Fixture que crea un directorio temporal."""
    with tempfile.TemporaryDirectory() as temp_dir:
        yield temp_dir


@pytest.fixture
def sample_docx_file(temp_dir):
    """Fixture que crea un archivo .docx de prueba."""
    from docx import Document
    
    # Crear documento de prueba
    doc = Document()
    doc.add_heading('Documento de Prueba', 0)
    doc.add_paragraph('Este es un párrafo de prueba.')
    doc.add_paragraph('Este es otro párrafo con texto en negrita.').bold = True
    
    # Agregar tabla
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Celda 1'
    table.cell(0, 1).text = 'Celda 2'
    table.cell(1, 0).text = 'Celda 3'
    table.cell(1, 1).text = 'Celda 4'
    
    # Guardar archivo
    file_path = Path(temp_dir) / "test_document.docx"
    doc.save(str(file_path))
    
    return str(file_path)


@pytest.fixture
def sample_doc_file(temp_dir):
    """Fixture que crea un archivo .doc de prueba (simulado)."""
    # Crear un archivo .doc simulado
    file_path = Path(temp_dir) / "test_document.doc"
    file_path.write_text("Contenido simulado de archivo .doc")
    
    return str(file_path)


@pytest.fixture
def invalid_file(temp_dir):
    """Fixture que crea un archivo inválido."""
    file_path = Path(temp_dir) / "invalid_file.txt"
    file_path.write_text("Este no es un archivo Word válido")
    
    return str(file_path)


@pytest.fixture
def large_file(temp_dir):
    """Fixture que crea un archivo grande (simulado)."""
    file_path = Path(temp_dir) / "large_file.docx"
    
    # Crear contenido grande (simulado)
    large_content = "x" * (200 * 1024 * 1024)  # 200MB
    file_path.write_text(large_content)
    
    return str(file_path)


@pytest.fixture
def mock_config():
    """Fixture que mockea la configuración."""
    with patch('src.config.config') as mock_config:
        mock_config.max_file_size = 104857600  # 100MB
        mock_config.supported_extensions = [".doc", ".docx"]
        mock_config.max_workers = 4
        mock_config.timeout = 300
        mock_config.log_level = "INFO"
        mock_config.default_output_dir = "./output"
        mock_config.output_quality = "high"
        yield mock_config


@pytest.fixture
def mock_logger():
    """Fixture que mockea el logger."""
    with patch('src.utils.logger.get_logger') as mock_logger:
        mock_logger.return_value = Mock()
        yield mock_logger


@pytest.fixture
def sample_files_list(temp_dir):
    """Fixture que crea una lista de archivos de prueba."""
    files = []
    
    # Crear varios archivos .docx
    for i in range(5):
        from docx import Document
        doc = Document()
        doc.add_heading(f'Documento {i+1}', 0)
        doc.add_paragraph(f'Contenido del documento {i+1}')
        
        file_path = Path(temp_dir) / f"document_{i+1}.docx"
        doc.save(str(file_path))
        files.append(str(file_path))
    
    return files


@pytest.fixture
def output_directory(temp_dir):
    """Fixture que crea un directorio de salida."""
    output_dir = Path(temp_dir) / "output"
    output_dir.mkdir(exist_ok=True)
    return str(output_dir)


@pytest.fixture
def mock_magic():
    """Fixture que mockea la librería magic."""
    with patch('src.validators.file_validator.magic') as mock_magic:
        mock_magic.from_file.return_value = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        yield mock_magic 